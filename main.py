import zipfile
import os
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import tempfile
# from PIL import Image

def insert_image_into_docx(input_docx, image_path, output_docx, paragraph_count, width, height):
    """
    Inserts an image into a specific paragraph of a DOCX file and resizes it.

    Args:
        input_docx (str): Path to the input DOCX file.
        image_path (str): Path to the image file to be inserted.
        output_docx (str): Path to save the modified DOCX file.
        paragraph_count (int): The 1-based index of the paragraph where the image should be inserted.
        width (float): Desired width of the inserted image (in inches).
        height (float): Desired height of the inserted image (in inches).
    """
    # Create a new Document object
    doc = Document(input_docx)

    # Ensure that the paragraph_count is within a valid range
    if paragraph_count < 1 or paragraph_count > len(doc.paragraphs):
        raise ValueError("Invalid paragraph_count. It should be within the range of paragraphs in the document.")

    # Calculate the position in points (72 points = 1 inch)
    width_points = width * 72
    height_points = height * 72

    # Get the specified paragraph
    target_paragraph = doc.paragraphs[paragraph_count - 1]

    # Add a run to the paragraph to contain the image
    run = target_paragraph.add_run()

    # Add the image to the run
    pic = run.add_picture(image_path, width=Inches(width), height=Inches(height))

    # Save the modified document
    doc.save(output_docx)

    print(f"Image inserted into paragraph {paragraph_count} and resized to width={width} inches, height={height} inches, saved to {output_docx}")

# Example usage:

# insert_image_into_docx("thank.docx", "sig.jpg", "output.docx", left=100, top=100, width=0.6, height=0.6)
def count_paragraphs_and_print_words(docx_file):
    """
    Count paragraphs in a DOCX file and print the words in each paragraph.

    Args:
        docx_file (str): Path to the input DOCX file.
    """
    # Create a Document object from the DOCX file
    doc = Document(docx_file)

    # Iterate through paragraphs and print the words
    # for paragraph in doc.paragraphs:
    #     paragraph_text = paragraph.text
    #     word_count = len(paragraph_text.split())
    #     print(f"Paragraph: {paragraph_text}")
    #     print(f"Word Count: {word_count}")
    #     print("-" * 20)

    # # Print the total number of paragraphs
    # total_paragraphs = len(doc.paragraphs)
    # print(f"Total Paragraphs: {total_paragraphs}")
    sig_loc = 0
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        paragraph_text = paragraph.text
        word_count = len(paragraph_text.split())
        print(f"Paragraph {i}:")
        print(f"Text: {paragraph_text}")
        if paragraph.text == "Adewale, Adedotun and Olufunso Odugbesan.":
            # print("Yesssssss!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            sig_loc = i
        print(f"Word Count: {word_count}")
        print("-" * 20)

    # Print the total number of paragraphs
    total_paragraphs = len(doc.paragraphs)
    print(f"Total Paragraphs: {total_paragraphs}")
    return sig_loc


def unzip_file(zip_file, output_dir):
    try:
        os.makedirs(output_dir, exist_ok=True)

        with zipfile.ZipFile(zip_file, 'r') as zip_ref:
            zip_ref.extractall(output_dir)
        
        print(f"Unzipped {unzip_file} to {output_dir}")
    except Exception as e:
        print(f"An errror occured: {e}")




def process_folder(folder_path, image_path):
    # Iterate through files in the folder
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            docx_file = os.path.join(folder_path, filename)
            sig_loc = count_paragraphs_and_print_words(docx_file)
            print(sig_loc)
            if sig_loc:
                # Modify the output file name as needed
                output_docx = os.path.join(folder_path, f"{filename}")
                insert_image_into_docx(docx_file, image_path, output_docx, paragraph_count=sig_loc - 1, width=0.5, height=0.5)
# Example usage:
# docx_file = "uncle.docx"  # Replace with your DOCX file path
# sig_loc = count_paragraphs_and_print_words(docx_file)
# print(sig_loc)
# insert_image_into_docx("uncle.docx", "sig.jpg", "output.docx", paragraph_count=int(sig_loc) - 1, width=0.5, height=0.5)


output_folder = "all_signed"
image_path = "sig.jpg"

# process_folder(folder_path, image_path)
app = FastAPI()
origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
@app.post("/upload-zip")
async def upload_zip_files(file1: UploadFile, file2: UploadFile, signature_image: UploadFile = File(...)):
    try:
        # Save the uploaded zip files temporarily
        zip_file1_path = os.path.join(".", file1.filename)
        zip_file2_path = os.path.join(".", file2.filename)
        
        with open(zip_file1_path, "wb") as f1, open(zip_file2_path, "wb") as f2:
            f1.write(file1.file.read())
            f2.write(file2.file.read())
        
        # Save the uploaded signature image
        signature_image_path = os.path.join(".", signature_image.filename)
        with open(signature_image_path, "wb") as sig_file:
            sig_file.write(signature_image.file.read())
        
        # Unzip the files
        for zip_file_path in [zip_file1_path, zip_file2_path]:
            with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
                zip_ref.extractall(output_folder)
        process_folder("all_signed", signature_image_path)
        # # Call the process_folder function for each unzipped folder
        # for folder_name in os.listdir(output_folder):
        #     folder_path = os.path.join(output_folder, folder_name)
        #     if os.path.isdir(folder_path):
        #         # process_folder(folder_path, signature_image_path)
        #         process_folder(folder_path, "sig.jpg")

        return JSONResponse(content={"message": "Files unzipped and processed successfully"})
    except Exception as e:
        return JSONResponse(content={"message": f"An error occurred: {str(e)}"}, status_code=500)

@app.get("/download/{folder_name}")
async def download_folder(folder_name: str):
    try:
        folder_path = os.path.join("./", folder_name)

        if not os.path.exists(folder_path):
            raise HTTPException(status_code=404, detail="Folder not found")

        # Create a temporary directory to store the zip archive
        temp_dir = tempfile.mkdtemp()

        # Create a zip file to store the folder contents
        zip_filename = f"{folder_name}.zip"
        zip_filepath = os.path.join(temp_dir, zip_filename)

        with zipfile.ZipFile(zip_filepath, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zipf.write(file_path, arcname=arcname)

        # Serve the zip archive for download
        return FileResponse(zip_filepath, media_type='application/zip', filename=zip_filename)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


if __name__ == "__main__":
    import uvicorn 
    uvicorn.run(app, host="0.0.0.0", port=8002)
    

    


